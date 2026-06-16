#!/usr/bin/env python3
"""
Build a .docx from a .tex manuscript, matching the original python-docx
pipeline used for manuscript.docx / manuscript_supplementary.docx.

Two-stage: pandoc converts .tex → plain text (handling all LaTeX complexity),
then python-docx assembles the .docx with embedded figures.

Usage:
    python build_docx_from_tex.py                            # m.tex→m.docx, m_s.tex→m_s.docx
    python build_docx_from_tex.py -i manuscript.tex -o manuscript.docx
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent

# Figure key → 300-DPI PNG (pre-generated in repo)
FIG_KEYS = {
    "fig1_mechanism": "experiments/results_v6/fig1_mechanism.docx-300dpi.png",
    "fig2_quant": "experiments/results_v6/fig2_quant.docx-300dpi.png",
    "fig3_comparisons": "experiments/results_v6/fig3_comparisons.docx-300dpi.png",
    "figS1_tasks": "experiments/results_v6/figS1_tasks.docx-300dpi.png",
    "figS2_tasks_3d": "experiments/results_v6/figS2_tasks_3d.docx-300dpi.png",
}

CONVERSIONS = [
    ("m.tex", "m.docx"),
    ("m_s.tex", "m_s.docx"),
]


def find_pandoc() -> str:
    path = shutil.which("pandoc")
    if path:
        return path
    try:
        import pypandoc
        return pypandoc.get_pandoc_path()
    except Exception:
        print("pandoc not found. Install pandoc or `pip install pypandoc`.")
        sys.exit(1)


def tex_to_plaintext(tex_path: Path, pandoc_bin: str) -> str:
    """Use pandoc to convert .tex to clean plain text."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        result = subprocess.run(
            [
                pandoc_bin,
                str(tex_path),
                "-f", "latex",
                "-t", "plain",
                "--wrap=none",
                "-o", str(tmp_path),
            ],
            cwd=ROOT,
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            print(f"  pandoc warning: {result.stderr[-300:]}", file=sys.stderr)
        text = tmp_path.read_text(encoding="utf-8", errors="replace")
        return text
    finally:
        try:
            tmp_path.unlink()
        except Exception:
            pass


def _find_figures_in_tex(raw: str) -> dict[str, int]:
    """
    Scan raw .tex for \includegraphics commands.
    Returns {fig_key: line_index_in_raw}.
    We use line index to approximate where each figure should appear.
    """
    fig_lines: dict[str, int] = {}
    for i, line in enumerate(raw.splitlines()):
        m = re.search(r"\\includegraphics\[[^]]*\]\{([^}]*)\}", line)
        if m:
            path = m.group(1)
            stem = Path(path).stem
            for key in FIG_KEYS:
                if key in stem or stem in key:
                    fig_lines[key] = i
                    break
    return fig_lines


def add_figure(doc: Document, fig_key: str) -> None:
    """Embed a 300-DPI PNG into the document."""
    png_rel = FIG_KEYS.get(fig_key)
    if png_rel is None:
        return
    png_path = ROOT / png_rel
    if not png_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure: {fig_key}]")
        run.font.size = Pt(9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(5.5))


def _fig_key_from_plain_text(para_text: str) -> str | None:
    """Detect if a plain-text paragraph describes a known figure."""
    for key in FIG_KEYS:
        # e.g. "Figure 1" → "fig1" but we use direct key hints
        if key in para_text.lower():
            return key
    # Check by number pattern: "Figure 1." → fig1_mechanism, "Figure 2." → fig2_quant, etc.
    m = re.match(r"Figure\s+(\d+)", para_text, re.IGNORECASE)
    if m:
        num = int(m.group(1))
        mapping = {1: "fig1_mechanism", 2: "fig2_quant", 3: "fig3_comparisons"}
        return mapping.get(num)
    # Supplementary: "Figure S1" → figS1_tasks, "Figure S2" → figS2_tasks_3d
    m = re.match(r"Figure\s+S(\d+)", para_text, re.IGNORECASE)
    if m:
        num = int(m.group(1))
        mapping = {1: "figS1_tasks", 2: "figS2_tasks_3d"}
        return mapping.get(num)
    return None


def _extract_braced(raw: str, cmd: str) -> str | None:
    """Extract \\cmd{...} content with proper brace matching."""
    prefix = "\\" + cmd + "{"
    idx = raw.find(prefix)
    if idx < 0:
        return None
    start = idx + len(prefix)
    depth = 1
    i = start
    while i < len(raw) and depth > 0:
        if raw[i] == '{':
            depth += 1
        elif raw[i] == '}':
            depth -= 1
            if depth == 0:
                return raw[start:i]
        i += 1
    return None


def _unwrap_braces(text: str, cmd: str) -> str:
    """Remove \\cmd{...} wrapper, keeping inner content (handles nested braces)."""
    prefix = "\\" + cmd + "{"
    i = 0
    result = []
    while i < len(text):
        if text[i:i + len(prefix)] == prefix:
            depth = 1
            j = i + len(prefix)
            while j < len(text) and depth > 0:
                if text[j] == '{':
                    depth += 1
                elif text[j] == '}':
                    depth -= 1
                    if depth == 0:
                        break
                j += 1
            if depth == 0:
                inner = text[i + len(prefix):j]
                result.append(inner)
                i = j + 1
            else:
                result.append(text[i])
                i += 1
        else:
            result.append(text[i])
            i += 1
    return ''.join(result)


def _strip_braced_inline(s: str) -> str:
    """Strip inline LaTeX formatting commands from a string."""
    for cmd in [r"textbf", r"textit", r"emph", r"textsc", r"texttt",
                r"mathrm", r"mathbf", r"textrm", r"noindent"]:
        s = _unwrap_braces(s, cmd)
    s = s.replace(r"\noindent", "").replace(r"\\", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def tex_to_docx(tex_path: Path, docx_path: Path, pandoc_bin: str) -> None:
    raw = tex_path.read_text(encoding="utf-8", errors="replace")
    fig_line_map = _find_figures_in_tex(raw)

    # ---- Extract title and abstract from raw .tex ----
    title_text = _extract_braced(raw, "title")
    if title_text:
        title_text = _strip_braced_inline(title_text)

    abstract_text = None
    abs_start = raw.find(r"\begin{abstract}")
    abs_end = raw.find(r"\end{abstract}")
    if abs_start >= 0 and abs_end >= 0:
        # Get content inside abstract environment
        abs_inner = raw[abs_start + len(r"\begin{abstract}"):abs_end]
        abs_inner = _strip_braced_inline(abs_inner)
        abstract_text = abs_inner.strip()

    # ---- Stage 1: pandoc converts LaTeX → plain text ----
    plain = tex_to_plaintext(tex_path, pandoc_bin)

    # ---- Stage 2: python-docx assembles the .docx ----
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Add title
    if title_text:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)

    # Add abstract
    if abstract_text:
        p = doc.add_paragraph()
        run = p.add_run("Abstract")
        run.bold = True
        doc.add_paragraph(abstract_text)

    paragraphs = [p.strip() for p in plain.split("\n\n") if p.strip()]

    # Track figure insertion
    figs_inserted: set[str] = set()
    raw_lines = raw.splitlines()
    current_raw_line = 0

    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)

        # Detect heading patterns for bolding
        if re.match(r"^\d+(\.\d+)*\.?\s+\w", para_text):
            for run in p.runs:
                run.bold = True

        # Estimate position in raw .tex
        search = para_text[:60].strip()
        for i in range(current_raw_line, len(raw_lines)):
            if search and search[:30] in raw_lines[i]:
                current_raw_line = i
                break

        # Insert figures found in THIS file only
        for fig_key, fig_line in sorted(fig_line_map.items(), key=lambda x: x[1]):
            if fig_key in figs_inserted:
                continue
            if fig_line <= current_raw_line + 5:
                add_figure(doc, fig_key)
                figs_inserted.add(fig_key)

    # Insert only the remaining figures THAT WERE DETECTED in this file
    for fig_key in fig_line_map:
        if fig_key not in figs_inserted:
            add_figure(doc, fig_key)

    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"  {tex_path.name} → {docx_path.name}  ({size_kb:.0f} KB)")


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert .tex to .docx (pandoc + python-docx)")
    parser.add_argument("--input", "-i", help="Input .tex file")
    parser.add_argument("--output", "-o", help="Output .docx file")
    args = parser.parse_args()

    pandoc_bin = find_pandoc()
    print(f"pandoc: {pandoc_bin}")

    if args.input:
        tex = Path(args.input)
        if not tex.is_absolute():
            tex = ROOT / tex
        docx = Path(args.output) if args.output else tex.with_suffix(".docx")
        if not docx.is_absolute():
            docx = ROOT / docx
        tex_to_docx(tex, docx, pandoc_bin)
        return 0

    for tex_name, docx_name in CONVERSIONS:
        tex = ROOT / tex_name
        docx = ROOT / docx_name
        if tex.exists():
            tex_to_docx(tex, docx, pandoc_bin)
        else:
            print(f"SKIP: {tex} not found")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
