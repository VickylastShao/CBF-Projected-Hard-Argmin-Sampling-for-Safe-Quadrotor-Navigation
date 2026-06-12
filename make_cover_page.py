"""Generate IJRA cover page .docx with author information.
Required because the main manuscript is anonymised for double-blind review.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Wide-Scale Sampling with Hard-Argmin Selection and CBF Projection "
    "for Safe Quadrotor Navigation in Narrow Passages"
)
run.bold = True
run.font.size = Pt(14)

# --- Cover page label ---
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Cover Page — Author Information").italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Submitted to: ").bold = True
p.add_run(
    "International Journal of Robotics and Automation (IJRA), ACTA Press"
)
p = doc.add_paragraph()
p.add_run("Manuscript type: ").bold = True
p.add_run("Regular Paper")
p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("June 2026")

# --- Abstract ---
doc.add_paragraph()
doc.add_heading("Abstract", level=2)
abstract_text = (
    "A common implementation of sampling-based NMPC for quadrotor obstacle "
    "avoidance draws candidate commands from a narrow Gaussian around a PD "
    "nominal. In tight passages, this nominal can point toward the conservative "
    "composite-barrier boundary rather than through the actual geometric gap. "
    "Adding more samples close to that nominal does not address this failure "
    "mode --- it is a candidate-generation problem, not a solver problem. "
    "TSH-NMPC widens the sampling distribution (sigma = 5 N around the PD "
    "nominal at K = 10), selects the lowest-cost rollout by hard argmin, and "
    "projects the result through a DT-CBF quadratic program before execution. "
    "On the narrow-passage benchmark, the wide-scale sampler reached 97-99% "
    "success, while the narrow-scale PD-Gaussian baseline reached 64-70%. "
    "Under +50% mass mismatch the baseline failed in all 40 trials; the "
    "wide-scale variants retained 77.5%-82.5% success. Compared with "
    "CasADi+IPOPT, TSH-NMPC matched the binary success rate at roughly 6 ms "
    "per call with a fixed O(K*S) rollout budget, although CasADi retained "
    "lower terminal error. A learned-proposal ablation reduced success from "
    "97.7% to 88.7%, suggesting that coverage mattered more than prediction "
    "accuracy in this setting."
)
p = doc.add_paragraph(abstract_text)

# --- Authors ---
doc.add_paragraph()
doc.add_heading("Authors", level=2)

authors = [
    ("Zhuang Shao", "1", "shaozhuang@crpower.com.cn", True, "0000-0003-2496-0797"),
    ("Lijun Lei", "2", "leilijun6@crpower.com.cn", False, ""),
    ("Peng Wang", "3", "wangpeng@ncwu.edu.cn", False, ""),
    ("Liang Zheng", "2", "zhengliang35@crpower.com.cn", False, ""),
    ("Jie Zhou", "2", "zhoujie365@crpower.com.cn", False, ""),
]

for (name, aff, email, is_corr, orcid) in authors:
    p = doc.add_paragraph()
    mark = "*" if is_corr else ""
    run = p.add_run(name + mark)
    run.bold = True
    p.add_run(f"  [affiliation {aff}]").italic = True
    sub = doc.add_paragraph(style="List Bullet")
    sub.add_run(f"Email: {email}")
    if orcid:
        sub2 = doc.add_paragraph(style="List Bullet")
        sub2.add_run(f"ORCID: {orcid}")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("* Corresponding author").italic = True

# --- Affiliations ---
doc.add_paragraph()
doc.add_heading("Affiliations", level=2)

affiliations = [
    ("[1]",
     "China Resources Power Technology Research Institute Co., Ltd., "
     "Shenzhen 518000, Guangdong Province, China"),
    ("[2]",
     "Rundian Energy Science and Technology Co., Ltd., "
     "Zhengzhou 450052, Henan Province, China"),
    ("[3]",
     "North China University of Water Resources and Electric Power, "
     "Zhengzhou 450046, Henan Province, China"),
]

for label, desc in affiliations:
    p = doc.add_paragraph()
    p.add_run(f"{label} ").bold = True
    p.add_run(desc)

# --- Declarations ---
doc.add_paragraph()
doc.add_heading("Declarations", level=2)

p = doc.add_paragraph()
p.add_run("Conflict of interest: ").bold = True
p.add_run("The authors declare no conflict of interest.")

p = doc.add_paragraph()
p.add_run("Funding: ").bold = True
p.add_run(
    "This work was supported by an internal research program of "
    "China Resources Power Technology Research Institute Co., Ltd."
)

p = doc.add_paragraph()
p.add_run("Data availability: ").bold = True
p.add_run(
    "The experimental logs, plotting scripts, and implementation code "
    "are publicly available at "
    "https://github.com/VickylastShao/"
    "CBF-Projected-Hard-Argmin-Sampling-for-Safe-Quadrotor-Navigation"
)

p = doc.add_paragraph()
p.add_run("Author contributions: ").bold = True
p.add_run(
    "Z. Shao: Conceptualization, Methodology, Software, Formal analysis, "
    "Writing --- original draft. L. Lei: Validation, Experimental design, "
    "Writing --- review and editing. P. Wang: Validation, Writing --- review "
    "and editing. L. Zheng: Visualization, Data curation. J. Zhou: "
    "Supervision, Project administration, Writing --- review and editing. "
    "All authors approved the final manuscript."
)

# --- Suggested Reviewers ---
doc.add_page_break()
doc.add_heading("Suggested Reviewers", level=2)

doc.add_paragraph(
    "The following researchers have substantial expertise in nonlinear "
    "control, model predictive control, and autonomous systems, and have "
    "not collaborated with any author of this manuscript within the past "
    "three years."
)

reviewers = [
    ("Prof. Jianxin Zhou",
     "zjx@seu.edu.cn",
     "School of Automation, Southeast University, Nanjing, China",
     "Nonlinear control, predictive control, unmanned aerial vehicle systems"),
    ("Prof. Zhenlong Wu",
     "wuzhenlong2020@zzu.edu.cn",
     "School of Electrical and Information Engineering, "
     "Zhengzhou University, Zhengzhou, China",
     "Robust control, model predictive control, industrial automation"),
    ("Prof. Cong Yu",
     "congy@jhun.edu.cn",
     "Jianghan University, Wuhan, China",
     "Control theory, nonlinear systems, robotics"),
    ("Prof. Hui Gu",
     "guhuini@126.com",
     "School of Automation, Nanjing University of Science and Technology, "
     "Nanjing, China",
     "Model predictive control, motion planning, autonomous systems"),
]

for i, (name, email, aff, expertise) in enumerate(reviewers):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"{i+1}. {name}").bold = True
    sub = doc.add_paragraph(style="List Bullet")
    sub.add_run(f"Email: {email}")
    sub = doc.add_paragraph(style="List Bullet")
    sub.add_run(f"Affiliation: {aff}")
    sub = doc.add_paragraph(style="List Bullet")
    sub.add_run(f"Expertise: {expertise}")

doc.save("cover_page_ijra.docx")
print("cover_page_ijra.docx written.")
