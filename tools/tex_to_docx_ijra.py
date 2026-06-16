#!/usr/bin/env python3
"""
将 IJRA 投稿稿（LaTeX）转换为 IJRA 要求的 Word .docx：
  - 12pt Times New Roman
  - Double-spaced
  - One column
  - 已匿名（不写入作者）
  - 公式 → Word 原生 OMML（latex2mathml + mathml2omml）
  - 表格 / itemize / enumerate / 引用 / 参考文献 完整保留

借鉴 super-patents/tools/md_to_docx.py 的公式管线，但直接处理 LaTeX
源文件（避开 markdown 中间格式对公式特殊字符的破坏）。

用法：
    python3 tools/tex_to_docx_ijra.py \\
        --input PTRM_NMPC_manuscript_ijra.tex \\
        --output PTRM_NMPC_manuscript_ijra.docx
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

# 复用 super-patents 的公式转换核心
sys.path.insert(0, str(Path("/mnt/c/Users/vicks/MyWork/super-patents/tools").resolve()))
from md_to_docx import (  # type: ignore
    _latex_to_omath,
    _preprocess_latex,
)

# ─── 全局 IJRA 格式 ───────────────────────────────────────────────
DEFAULT_FONT = "Times New Roman"
BODY_PT = 12.0
LINE_SPACING = 2.0  # double-spaced
SPACE_AFTER_PT = 0.0
HEADING_PT = 12.0
HEADING_BOLD = True
FIRST_LINE_INDENT_CM = 0.75  # 正文每段首行缩进（≈ ½ inch）

# ─── LaTeX 宏定义（与 ijra.tex 头部一致）────────────────────────
LATEX_MACROS = {
    r"\R": r"\mathbb{R}",
    r"\E": r"\mathbb{E}",
    r"\Nm": r"N_{\mathrm{MC}}",
    r"\spt": r"\mathrm{sp}",
    r"\pd": r"\mathrm{pd}",
    r"\dt": r"\Delta t",
    r"\nom": r"\mathrm{nom}",
    r"\safe": r"\mathrm{safe}",
    r"\TErr": r"\mathrm{TErr}",
    r"\IAE": r"\mathrm{IAE}",
}

# 计数器：用于 \label / \ref 解析
_SECTION_COUNTERS = {"section": 0, "subsection": 0, "subsubsection": 0}
_label_map: dict[str, str] = {}  # label_id → display_number
_equation_counter = 0
_table_counter = 0
_figure_counter = 0
_S_PREFIX = False  # set to True when processing supplementary files
_proposition_counter = 0
# 参考文献：按正文 \cite 顺序编号；未被引用的 bibitem 顺延列入末尾。
_cite_order: dict[str, int] = {}   # key → 数字编号
_cite_seq: list[str] = []          # 按引用先后排序的 key 列表


# ─── 文本处理工具 ─────────────────────────────────────────────────


def _strip_comments(s: str) -> str:
    """删除 LaTeX 注释 (% 起始至行尾，但保留 \\%)。"""
    out_lines = []
    for line in s.split("\n"):
        # 找未转义的 %
        i = 0
        result = []
        while i < len(line):
            if line[i] == "\\" and i + 1 < len(line):
                result.append(line[i : i + 2])
                i += 2
                continue
            if line[i] == "%":
                break
            result.append(line[i])
            i += 1
        out_lines.append("".join(result))
    return "\n".join(out_lines)


def _expand_macros(s: str) -> str:
    """展开 ijra.tex 自定义宏。按长名→短名顺序避免 \\Nm 被 \\N 替换。"""
    for macro in sorted(LATEX_MACROS.keys(), key=len, reverse=True):
        # 仅当后面跟非字母或 {} 或字符串结尾时替换
        pattern = re.escape(macro) + r"(?![A-Za-z])"
        s = re.sub(pattern, lambda _m: LATEX_MACROS[macro] + " ", s)
    return s


def _balanced_arg(s: str, start: int) -> Optional[tuple[str, int]]:
    """从 s[start]（应为 '{'）提取平衡花括号内容。返回 (内容, 结束位置)。"""
    if start >= len(s) or s[start] != "{":
        return None
    depth = 0
    i = start
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s):
            i += 2
            continue
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start + 1 : i], i + 1
        i += 1
    return None


# ─── 标签 / 引用预扫描 ───────────────────────────────────────────


def _prescan_labels(source: str) -> None:
    """两遍解析：第一遍扫描 \\label 建立 label → 编号 映射。"""
    sec = sub = subsub = 0
    eq_num = 0
    tab_num = 0
    fig_num = 0
    prop_num = 0
    alg_num = 0
    # 状态机：跟踪所在环境
    lines = source.split("\n")
    in_table_env = False
    cur_table_label: Optional[str] = None
    in_figure_env = False
    cur_figure_label: Optional[str] = None
    in_eq_env = False
    cur_eq_label: Optional[str] = None
    in_align_env = False
    align_label_count = 0
    in_prop = False
    cur_prop_label: Optional[str] = None
    in_alg_env = False
    cur_alg_label: Optional[str] = None

    i = 0
    while i < len(lines):
        line = lines[i]

        # 段落标题（使用平衡花括号，因为标题可能含 $...$）
        section_kind = None
        for cmd, key in ((r"\section", "section"),
                         (r"\subsection", "subsection"),
                         (r"\subsubsection", "subsubsection")):
            pat = re.escape(cmd) + r"\*?\s*\{"
            m = re.match(pat, line)
            if m:
                brace_start = line.find("{", m.end() - 1)
                arg = _balanced_arg(line, brace_start)
                if arg:
                    section_kind = (key, arg[0], arg[1])
                    break
        if section_kind:
            key, title, end_pos = section_kind
            if key == "section":
                sec += 1; sub = 0; subsub = 0
                num = str(sec)
            elif key == "subsection":
                sub += 1; subsub = 0
                num = f"{sec}.{sub}"
            else:
                subsub += 1
                num = f"{sec}.{sub}.{subsub}"
            # 查找标题里的 \label{...}（同行或下一行）
            after = line[end_pos:]
            lab_m = re.search(r"\\label\{([^}]+)\}", after)
            if lab_m:
                _label_map[lab_m.group(1)] = num
            elif i + 1 < len(lines):
                lab_m2 = re.match(r"\s*\\label\{([^}]+)\}", lines[i + 1])
                if lab_m2:
                    _label_map[lab_m2.group(1)] = num
            i += 1
            continue

        # \paragraph{...}\label{...}：把 label 映射到当前最深 section
        para_m = re.search(
            r"\\paragraph\{[^}]+\}\\label\{([^}]+)\}", line
        )
        if para_m:
            cur_num = f"{sec}.{sub}"
            if subsub:
                cur_num = f"{sec}.{sub}.{subsub}"
            _label_map[para_m.group(1)] = cur_num

        # 环境
        if r"\begin{table}" in line:
            in_table_env = True
            cur_table_label = None
        if r"\begin{figure}" in line:
            in_figure_env = True
            cur_figure_label = None
        if r"\begin{equation}" in line:
            in_eq_env = True
            cur_eq_label = None
        if r"\begin{align}" in line:
            in_align_env = True
            align_label_count = 0
        if r"\begin{proposition}" in line:
            in_prop = True
            cur_prop_label = None
        if r"\begin{algorithm}" in line:
            in_alg_env = True
            cur_alg_label = None

        # 在环境内的 \label（包括 \phantomsection\label{...} 锚点）
        for lm in re.finditer(r"\\label\{([^}]+)\}", line):
            lab = lm.group(1)
            if in_table_env:
                cur_table_label = lab
            elif in_figure_env:
                cur_figure_label = lab
            elif in_eq_env:
                cur_eq_label = lab
            elif in_align_env:
                align_label_count += 1
                eq_num += 1
                _label_map[lab] = str(eq_num)
            elif in_prop:
                cur_prop_label = lab
            elif in_alg_env:
                cur_alg_label = lab
            elif r"\phantomsection" in line:
                cur_num = str(sec)
                if sub:
                    cur_num = f"{sec}.{sub}"
                if subsub:
                    cur_num = f"{sec}.{sub}.{subsub}"
                _label_map[lab] = cur_num

        # 环境结束
        if r"\end{table}" in line and in_table_env:
            tab_num += 1
            if cur_table_label:
                _label_map[cur_table_label] = str(tab_num)
            in_table_env = False
        if r"\end{figure}" in line and in_figure_env:
            fig_num += 1
            if cur_figure_label:
                _label_map[cur_figure_label] = str(fig_num)
            in_figure_env = False
        if r"\end{equation}" in line and in_eq_env:
            eq_num += 1
            if cur_eq_label:
                _label_map[cur_eq_label] = str(eq_num)
            in_eq_env = False
        if r"\end{align}" in line and in_align_env:
            in_align_env = False
        if r"\end{proposition}" in line and in_prop:
            prop_num += 1
            if cur_prop_label:
                _label_map[cur_prop_label] = str(prop_num)
            in_prop = False
        if r"\end{algorithm}" in line and in_alg_env:
            alg_num += 1
            if cur_alg_label:
                _label_map[cur_alg_label] = str(alg_num)
            in_alg_env = False
        i += 1


def _prescan_citations(source: str) -> None:
    """按正文 \\cite 出现顺序给 bibitem 编号；未被引用的 bibitem 顺延列入末尾。
    填充全局 _cite_order: {key → number} 和 _cite_seq: [key, ...]。
    """
    global _cite_order, _cite_seq
    _cite_order = {}
    _cite_seq = []
    # 找 thebibliography 起点，正文 = 之前的部分
    bib_start = source.find(r"\begin{thebibliography}")
    body = source[:bib_start] if bib_start >= 0 else source
    # 抓 \cite{a,b,c}（\cite[opt]{} 在 normalize 阶段已剥光可选参数）
    for m in re.finditer(r"\\cite\s*\{([^}]+)\}", body):
        for raw_key in m.group(1).split(","):
            key = raw_key.strip()
            if key and key not in _cite_order:
                _cite_seq.append(key)
                _cite_order[key] = len(_cite_seq)
    # bibitem 里未被引用的，追加在末尾（不丢条目）
    for m in re.finditer(r"\\bibitem\s*\{([^}]+)\}", source):
        key = m.group(1).strip()
        if key and key not in _cite_order:
            _cite_seq.append(key)
            _cite_order[key] = len(_cite_seq)


# ─── 行内文本 → docx run（处理 \cite, \ref, \emph, \textbf 等）─────


_INLINE_TOKEN_RE = re.compile(
    r"(\$\$[^$]+\$\$"  # display math (inline form)
    r"|\$[^$]+?\$"  # inline math
    r"|\\textbf\{[^}]*\}"
    r"|\\emph\{[^}]*\}"
    r"|\\textit\{[^}]*\}"
    r"|\\textsc\{[^}]*\}"
    r"|\\texttt\{[^}]*\}"
    r"|\\paragraph\{[^}]*\}"
    r"|\\cite\{[^}]+\}"
    r"|\\eqref\{[^}]+\}"
    r"|\\ref\{[^}]+\}"
    r"|\\label\{[^}]+\}"
    r"|\\url\{[^}]+\}"
    r"|\\href\{[^}]+\}\{[^}]*\}"
    r"|\\footnote\{[^}]*\}"
    r"|\\noindent\b"
    r"|\\medskip\b"
    r"|\\smallskip\b"
    r"|\\bigskip\b"
    r"|\\\\\b"
    # 重音命令需在普通 ` ' 引号 token 之前，优先成对捕获
    r"|\\[\"'`^~=]\{?[a-zA-Z]\}?"
    r"|\{\\[\"'`^~=][a-zA-Z]\}"
    r"|---"
    r"|--"
    r"|``|''|`|'"
    r"|~"
    r")"
)


def _decode_latex_escapes(s: str) -> str:
    """处理常见 LaTeX 字符转义和重音命令（多用于参考文献）。"""
    # 重音：\"o → ö, \'a → á, \`a → à, \^o → ô, \~n → ñ, \=a → ā
    accents = {
        '"': {'a': 'ä', 'A': 'Ä', 'o': 'ö', 'O': 'Ö', 'u': 'ü', 'U': 'Ü',
              'e': 'ë', 'E': 'Ë', 'i': 'ï', 'I': 'Ï', 'y': 'ÿ'},
        "'": {'a': 'á', 'A': 'Á', 'e': 'é', 'E': 'É', 'i': 'í', 'I': 'Í',
              'o': 'ó', 'O': 'Ó', 'u': 'ú', 'U': 'Ú', 'y': 'ý', 'n': 'ń',
              'c': 'ć', 's': 'ś', 'z': 'ź'},
        '`': {'a': 'à', 'A': 'À', 'e': 'è', 'E': 'È', 'i': 'ì', 'I': 'Ì',
              'o': 'ò', 'O': 'Ò', 'u': 'ù', 'U': 'Ù'},
        '^': {'a': 'â', 'A': 'Â', 'e': 'ê', 'E': 'Ê', 'i': 'î', 'I': 'Î',
              'o': 'ô', 'O': 'Ô', 'u': 'û', 'U': 'Û'},
        '~': {'a': 'ã', 'A': 'Ã', 'n': 'ñ', 'N': 'Ñ', 'o': 'õ', 'O': 'Õ'},
        '=': {'a': 'ā', 'e': 'ē', 'i': 'ī', 'o': 'ō', 'u': 'ū'},
    }
    def acc_sub(m):
        mark, letter = m.group(1), m.group(2)
        return accents.get(mark, {}).get(letter, m.group(0))
    # 顺序很关键：先 {\"o} 形式（防止 \{ 被 \\{ 替换吞掉），再 \"{o}，
    # 最后 \"o 形式
    s = re.sub(r'\{\\(["\'`^~=])([a-zA-Z])\}', acc_sub, s)
    s = re.sub(r'\\(["\'`^~=])\{([a-zA-Z])\}', acc_sub, s)
    s = re.sub(r'\\(["\'`^~=])([a-zA-Z])', acc_sub, s)
    # 特殊连字 / 字母
    s = s.replace(r"\ss", "ß")
    s = s.replace(r"\AE", "Æ").replace(r"\ae", "æ")
    s = s.replace(r"\OE", "Œ").replace(r"\oe", "œ")
    s = s.replace(r"\AA", "Å").replace(r"\aa", "å")
    s = s.replace(r"\O", "Ø").replace(r"\o", "ø")
    s = s.replace(r"\Large", "").replace(r"\large", "")
    s = s.replace(r"\L", "Ł").replace(r"\l", "ł")
    # 常用字符转义
    s = s.replace(r"\&", "&")
    s = s.replace(r"\%", "%")
    s = s.replace(r"\#", "#")
    s = s.replace(r"\_", "_")
    s = s.replace(r"\$", "$")
    s = s.replace(r"\{", "{")
    s = s.replace(r"\}", "}")
    s = s.replace(r"\,", " ")
    s = s.replace(r"\;", " ")
    s = s.replace(r"\:", " ")
    s = s.replace(r"\!", "")
    s = s.replace(r"\ ", " ")
    s = s.replace("~", " ")
    s = s.replace("---", "—")
    s = s.replace("--", "–")
    s = s.replace("``", "“").replace("''", "”")
    # 不可见命令：\phantomsection（超链接锚点），孤立反斜杠（如 .strip() 残留）
    s = s.replace(r"\phantomsection", "")
    if s.endswith("\\"):
        s = s[:-1]
    return s


def _set_run_font(run, *, size_pt: float = BODY_PT, bold: bool = False,
                  italic: bool = False, mono: bool = False,
                  smallcaps: bool = False,
                  superscript: bool = False, subscript: bool = False,
                  color: Optional[RGBColor] = None):
    font_name = "Consolas" if mono else DEFAULT_FONT
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if smallcaps:
        # OOXML 小型大写
        sc = rpr.find(qn("w:smallCaps"))
        if sc is None:
            sc = etree.SubElement(rpr, qn("w:smallCaps"))
        sc.set(qn("w:val"), "true")
    if superscript:
        run.font.superscript = True
    if subscript:
        run.font.subscript = True
    if color is not None:
        run.font.color.rgb = color


def _format_paragraph(p, *, indent_cm: float = 0.0, space_after_pt: float = 0.0,
                      align=None, keep_with_next: bool = False,
                      first_line_indent_cm: float = 0.0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(0)
    if indent_cm:
        pf.left_indent = Cm(indent_cm)
    if first_line_indent_cm:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if align is not None:
        p.alignment = align
    if keep_with_next:
        pf.keep_with_next = True


# LaTeX 数学符号 → Unicode 映射（用于降级简单公式为纯文本）
_LATEX_UNICODE_MATH = {
    # 希腊字母
    r"\sigma": "σ", r"\Sigma": "Σ",
    r"\kappa": "κ", r"\alpha": "α", r"\gamma": "γ", r"\delta": "δ",
    r"\rho": "ρ", r"\tau": "τ", r"\epsilon": "ε", r"\mu": "μ",
    r"\beta": "β", r"\lambda": "λ", r"\omega": "ω", r"\Omega": "Ω",
    # 运算符 / 关系
    r"\ge": "≥", r"\le": "≤", r"\gg": "≫", r"\ll": "≪",
    r"\times": "×", r"\pm": "±", r"\mp": "∓",
    r"\approx": "≈", r"\sim": "∼", r"\simeq": "≃",
    r"\cdot": "·", r"\ldots": "…", r"\cdots": "⋯",
    r"\to": "→", r"\rightarrow": "→", r"\leftarrow": "←", r"\gets": "←",
    r"\top": "⊤", r"\bot": "⊥",
    r"\in": "∈", r"\notin": "∉", r"\subset": "⊂", r"\subseteq": "⊆",
    r"\cup": "∪", r"\cap": "∩",
    r"\infty": "∞", r"\partial": "∂",
    r"\forall": "∀", r"\exists": "∃",
    r"\rightarrow": "→", r"\Rightarrow": "⇒",
    r"\leftrightarrow": "↔",
    r"\min": "min", r"\max": "max", r"\arg": "arg",
    r"\log": "log", r"\exp": "exp", r"\Pr": "Pr",
    # 杂项
    r"\mathbb{R}": "ℝ", r"\mathbb{E}": "𝔼", r"\mathbb{N}": "ℕ",
    r"\mathcal{U}": "𝒰", r"\mathcal{N}": "𝒩",
    r"\text{s.t.}": "s.t.",
    r"\operatorname{clip}": "clip",
    r"\operatorname{diag}": "diag",
    r"\%": "%", r"\$": "$", r"\&": "&", r"\#": "#",
    r"\,": " ", r"\ ": " ", r"\;": " ", r"\:": " ", r"\!": "",
}

# 仅改变字体的 LaTeX 命令 —— 提取内部内容即可
_FONT_CMDS = ["mathrm", "text", "mathbf", "mathit", "mathsf", "mathtt",
              "textrm", "textit", "textbf"]


def _try_text_math(latex_str: str) -> Optional[str]:
    """若 latex 公式足够简单，返回 Unicode 纯文本；否则返回 None。"""
    s = latex_str.strip()
    if not s:
        return None
    # 含数组 / 分式 / 根式 / 求和积分极限 → OMML
    if re.search(r"\\begin\{|\\frac|\\sqrt|\\choose|\\binom|\\sum|\\prod|\\int|\\lim", s):
        return None
    # 去除 {,} 装饰（LaTeX 中防逗号间距的写法）
    s = s.replace("{,}", ",")
    # 展开字体命令 \cmd{inner} → inner（必须在去孤立花括号之前）
    for cmd in _FONT_CMDS:
        pattern = re.escape("\\" + cmd) + r"\s*\{(.*?)\}"
        s = re.sub(pattern, r"\1", s)
    # 去除仅用于分组/防断行的孤立花括号（不影响上下标检测，上下标需有 _ ^）
    s = re.sub(r"\{([^}]+)\}(?![_^])", r"\1", s)
    # 替换已知 LaTeX 命令为 Unicode（按长度降序避免部分匹配）
    for cmd in sorted(_LATEX_UNICODE_MATH.keys(), key=len, reverse=True):
        s = s.replace(cmd, _LATEX_UNICODE_MATH[cmd])
    # 若仍残留 \ 命令 → 交给 OMML
    if "\\" in s:
        return None
    # 含花括号 / 下标 / 上标 → 交给 OMML（数学渲染更规范）
    if re.search(r"[\{\}_\^]", s):
        return None
    # 清理多余空格
    result = re.sub(r" +", " ", s).strip()
    return result if result else None


def _add_inline_math(paragraph, latex_str: str) -> None:
    """行内公式 → OMML 或降级纯文本。"""
    expanded = _expand_macros(latex_str)
    # LaTeX 间距命令（\, \; \: \!）在 latex2mathml 中可能不识别，替换为普通空格
    for sp_cmd in [r'\,', r'\;', r'\:', r'\!']:
        expanded = expanded.replace(sp_cmd, ' ')
    # 公式末尾 LaTeX 标点（,/;/.）在 Word 公式里没有意义，
    # 反而会变成游离符号；这里剥掉。
    expanded = re.sub(r"\s*[,;]\s*$", "", expanded).strip()
    # 尝试降级为简单 Unicode 文本
    text_math = _try_text_math(expanded)
    if text_math is not None:
        # 按字母/符号 vs 数字/标点切分，字母部分斜体，数字部分正体
        parts = re.split(r'(\d+(?:\.\d+)?(?:[eE][+-]?\d+)?%?)', text_math)
        for part in parts:
            if not part:
                continue
            run = paragraph.add_run(part)
            is_num = re.match(r'^[\d\.\-–,\\%\s]+$', part)
            _set_run_font(run, size_pt=BODY_PT, italic=not is_num)
        return
    omath = _latex_to_omath(expanded)
    if omath is not None:
        paragraph._element.append(omath)
        return
    run = paragraph.add_run(latex_str)
    _set_run_font(run, size_pt=BODY_PT)


def _resolve_ref(key: str) -> str:
    val = _label_map.get(key, "?")
    if _S_PREFIX and val != "?":
        # 对于 supplementary，表格/图片/章节/算法编号加 S 前缀
        if key.startswith(("tab:", "fig:", "alg:")) or key.startswith("sec:S"):
            val = "S" + val
    return val


def _render_inline_segment(paragraph, text: str, *, bold: bool = False,
                           italic: bool = False) -> None:
    """渲染一段 LaTeX 内联文本到 paragraph，分割 token 并处理样式。"""
    if not text:
        return
    # 段内物理换行规整成空格；防止 python-docx 把 \n 渲染成软换行 <w:br/>。
    # 真正的段落分隔由上层 _split_logical_blocks 通过 \n\n 完成。
    if "\n" in text:
        text = re.sub(r"\s*\n\s*", " ", text)
        text = re.sub(r" {2,}", " ", text)
    # 先用平衡花括号扫描器把 \texttt{..} \textbf{..} \emph{..} \textit{..}
    # \paragraph{..} 这类参数中可能含 $..$ 或嵌套 { } 的命令切出来。
    # token 形态：('cmd', cmd_name, inner) 或 ('raw', text)
    BAL_CMDS = ("texttt", "textbf", "emph", "textit", "textsc", "paragraph",
                "mathrm", "mathbf")
    bal_re = re.compile(r"\\(" + "|".join(BAL_CMDS) + r")\s*\{")
    segments = []  # list of (kind, payload, opt_cmd)
    i = 0
    n = len(text)
    buf = []
    while i < n:
        # 优先跳过 $..$ math 段（math 内不解析文本命令）
        if text[i] == "$":
            j = text.find("$", i + 1)
            if j < 0:
                buf.append(text[i:])
                i = n
                break
            buf.append(text[i:j + 1])
            i = j + 1
            continue
        m = bal_re.match(text, i)
        if m:
            arg = _balanced_arg(text, m.end() - 1)
            if arg:
                inner, end_pos = arg
                if buf:
                    segments.append(("raw", "".join(buf), None))
                    buf = []
                segments.append(("cmd", inner, m.group(1)))
                i = end_pos
                continue
        buf.append(text[i])
        i += 1
    if buf:
        segments.append(("raw", "".join(buf), None))

    for kind, payload, cmd in segments:
        if kind == "cmd":
            if cmd == "textbf":
                _render_inline_segment(paragraph, payload, bold=True,
                                       italic=italic)
            elif cmd in ("emph", "textit"):
                _render_inline_segment(paragraph, payload, bold=bold,
                                       italic=True)
            elif cmd == "textsc":
                # \textsc{...} → 小型大写。Times New Roman 自带，
                # Word 用 w:smallCaps 渲染。
                inner = _decode_latex_escapes(payload)
                run = paragraph.add_run(inner)
                _set_run_font(run, size_pt=BODY_PT, bold=bold,
                              italic=italic, smallcaps=True)
            elif cmd == "texttt":
                # \texttt{...} 在投稿稿里多为策略代号/环境名（R1_s5、A3、
                # narrow、two_gate、PD 等）。为避免 Times 正文 + Consolas
                # mono + Cambria Math 三字体混排造成视觉割裂，统一渲染为
                # 与正文同字体同字号的 Times New Roman 12pt 正体。
                # 若 \texttt{} 内含 $..$ 数学，则把数学独立成 OMML，
                # 其余文本仍用正文字体。
                if "$" in payload:
                    sub_pos = 0
                    for mm in re.finditer(r"\$([^$]+)\$", payload):
                        if mm.start() > sub_pos:
                            plain = _decode_latex_escapes(
                                payload[sub_pos:mm.start()])
                            if plain:
                                run = paragraph.add_run(plain)
                                _set_run_font(run, size_pt=BODY_PT,
                                              bold=bold, italic=italic)
                        _add_inline_math(paragraph, mm.group(1).strip())
                        sub_pos = mm.end()
                    if sub_pos < len(payload):
                        plain = _decode_latex_escapes(payload[sub_pos:])
                        if plain:
                            run = paragraph.add_run(plain)
                            _set_run_font(run, size_pt=BODY_PT,
                                          bold=bold, italic=italic)
                else:
                    inner = _decode_latex_escapes(payload)
                    run = paragraph.add_run(inner)
                    _set_run_font(run, size_pt=BODY_PT, bold=bold,
                                  italic=italic)
            elif cmd == "paragraph":
                # \paragraph{...} 行内出现时作为粗体小标题
                _render_inline_segment(paragraph, payload + " ",
                                       bold=True, italic=italic)
            elif cmd in ("mathrm", "mathbf"):
                # 文本环境中的 \mathrm/\mathbf 直接取内部
                _render_inline_segment(paragraph, payload, bold=bold,
                                       italic=italic)
            continue
        # kind == 'raw'：走原有 token 流程
        _render_inline_raw(paragraph, payload, bold=bold, italic=italic)


def _render_inline_raw(paragraph, text: str, *, bold: bool = False,
                       italic: bool = False) -> None:
    """渲染不含平衡花括号文本命令的纯片段（math / 引用 / 转义等）。"""
    if not text:
        return
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            plain = _decode_latex_escapes(text[pos : m.start()])
            if plain:
                run = paragraph.add_run(plain)
                _set_run_font(run, bold=bold, italic=italic)
        tok = m.group(1)
        if tok.startswith("$$") and tok.endswith("$$"):
            _add_inline_math(paragraph, tok[2:-2].strip())
        elif tok.startswith("$") and tok.endswith("$"):
            _add_inline_math(paragraph, tok[1:-1].strip())
        elif tok.startswith(r"\textbf{"):
            inner = tok[len(r"\textbf{") : -1]
            _render_inline_segment(paragraph, inner, bold=True, italic=italic)
        elif tok.startswith(r"\emph{") or tok.startswith(r"\textit{"):
            prefix = r"\emph{" if tok.startswith(r"\emph{") else r"\textit{"
            inner = tok[len(prefix) : -1]
            _render_inline_segment(paragraph, inner, bold=bold, italic=True)
        elif tok.startswith(r"\textsc{"):
            inner = _decode_latex_escapes(tok[len(r"\textsc{") : -1])
            run = paragraph.add_run(inner)
            _set_run_font(run, size_pt=BODY_PT, bold=bold,
                          italic=italic, smallcaps=True)
        elif tok.startswith(r"\texttt{"):
            inner = _decode_latex_escapes(tok[len(r"\texttt{") : -1])
            run = paragraph.add_run(inner)
            _set_run_font(run, size_pt=BODY_PT, bold=bold, italic=italic)
        elif tok.startswith(r"\paragraph{"):
            inner = _decode_latex_escapes(tok[len(r"\paragraph{") : -1])
            run = paragraph.add_run(inner + " ")
            _set_run_font(run, bold=True, italic=italic)
        elif tok.startswith(r"\cite{"):
            keys = [k.strip() for k in tok[len(r"\cite{") : -1].split(",")]
            # 按 _cite_order 翻译成 ACTA 数字编号 [n]，
            # 连续编号合并为 [1,3,5] 但不合并为 [1-3]（ACTA 通用风格）
            nums = []
            for k in keys:
                if k in _cite_order:
                    nums.append(str(_cite_order[k]))
                else:
                    nums.append("?")
            run = paragraph.add_run("[" + ", ".join(nums) + "]")
            # ACTA / IJRA 风格：数字引用置于上标
            _set_run_font(run, bold=bold, italic=italic,
                          size_pt=BODY_PT, superscript=True)
        elif tok.startswith(r"\eqref{"):
            key = tok[len(r"\eqref{") : -1]
            run = paragraph.add_run("(" + _resolve_ref(key) + ")")
            _set_run_font(run, bold=bold, italic=italic)
        elif tok.startswith(r"\ref{"):
            key = tok[len(r"\ref{") : -1]
            run = paragraph.add_run(_resolve_ref(key))
            _set_run_font(run, bold=bold, italic=italic)
        elif tok.startswith(r"\label{"):
            pass  # 已预扫描
        elif tok.startswith(r"\url{"):
            inner = tok[len(r"\url{") : -1]
            run = paragraph.add_run(inner)
            _set_run_font(run, bold=bold, italic=italic, mono=True,
                          size_pt=BODY_PT - 1)
        elif tok.startswith(r"\href{"):
            m2 = re.match(r"\\href\{([^}]+)\}\{([^}]*)\}", tok)
            if m2:
                run = paragraph.add_run(m2.group(2) or m2.group(1))
                _set_run_font(run, bold=bold, italic=italic)
        elif tok.startswith(r"\footnote{"):
            inner = _decode_latex_escapes(tok[len(r"\footnote{") : -1])
            run = paragraph.add_run(f" [Note: {inner}]")
            _set_run_font(run, size_pt=BODY_PT - 2, italic=True,
                          color=RGBColor(0x55, 0x55, 0x55))
        elif tok in (r"\noindent", r"\medskip", r"\smallskip", r"\bigskip"):
            pass
        elif tok == r"\\":
            paragraph.add_run().add_break()
        elif tok == "---":
            run = paragraph.add_run("—")
            _set_run_font(run, bold=bold, italic=italic)
        elif tok == "--":
            run = paragraph.add_run("–")
            _set_run_font(run, bold=bold, italic=italic)
        elif tok == "~":
            # LaTeX ~ 是不间断空格；Word 中 NBSP 在某些字体下变占位圆圈，统一退化为普通空格。
            run = paragraph.add_run(" ")
            _set_run_font(run, bold=bold, italic=italic)
        elif tok in ("``", "''"):
            ch = "“" if tok == "``" else "”"
            run = paragraph.add_run(ch)
            _set_run_font(run, bold=bold, italic=italic)
        elif tok == "`":
            run = paragraph.add_run("‘")
            _set_run_font(run, bold=bold, italic=italic)
        elif tok == "'":
            run = paragraph.add_run("’")
            _set_run_font(run, bold=bold, italic=italic)
        elif re.match(r"\\[\"'`^~=]\{?[a-zA-Z]\}?", tok) or \
                re.match(r"\{\\[\"'`^~=][a-zA-Z]\}", tok):
            # 重音命令统一交给 _decode_latex_escapes 转 Unicode
            decoded = _decode_latex_escapes(tok)
            run = paragraph.add_run(decoded)
            _set_run_font(run, bold=bold, italic=italic)
        else:
            run = paragraph.add_run(tok)
            _set_run_font(run, bold=bold, italic=italic)
        pos = m.end()
    if pos < len(text):
        plain = _decode_latex_escapes(text[pos:])
        if plain:
            run = paragraph.add_run(plain)
            _set_run_font(run, bold=bold, italic=italic)


# ─── 主转换循环 ───────────────────────────────────────────────────


def _normalize_source(source: str) -> str:
    """去注释、展开宏、规整空白。"""
    s = _strip_comments(source)
    # \cite[\S IV]{KEY} → \cite{KEY}（丢弃可选参数）
    s = re.sub(r"\\cite\s*\[[^\]]*\]\s*\{", r"\\cite{", s)
    s = _expand_macros(s)
    # 兜底：所有 NBSP（U+00A0）与 NARROW NBSP（U+202F）以及 THIN SPACE（U+2009）
    # 一律转普通空格；防止在 Word 渲染为占位圆圈。
    s = s.replace("\u00a0", " ").replace("\u202f", " ").replace("\u2009", " ")
    return s


def _split_logical_blocks(source: str) -> list[tuple[str, str]]:
    """将正文（\\begin{document} ... \\end{document}）按段落/环境/标题切块。
    返回 [(kind, payload), ...]。
    """
    # 切出 document 主体
    m_doc = re.search(r"\\begin\{document\}(.*?)\\end\{document\}", source, re.DOTALL)
    if not m_doc:
        return []
    body = m_doc.group(1)

    blocks: list[tuple[str, str]] = []
    i = 0
    n = len(body)
    cur_para: list[str] = []
    # 当前 para 是否是"前一公式/对齐环境的延续段"——无空行紧跟在 \end{equation}
    # 之后的正文，按学术排版约定不应再加首行缩进。
    cont_flag = [False]

    def flush_para():
        if cur_para:
            text = "\n".join(cur_para).strip()
            if text:
                kind = "para_cont" if cont_flag[0] else "para"
                blocks.append((kind, text))
            cur_para.clear()
        cont_flag[0] = False

    # 我们逐行扫描，遇到环境时整段提取
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行 → 段落分隔
        if stripped == "":
            flush_para()
            i += 1
            continue

        # \title / \maketitle / \date 等元数据：直接处理
        m = re.match(r"\\title\s*\{(.*)", line)
        if m:
            # 多行标题
            title_buf = m.group(1)
            depth = title_buf.count("{") - title_buf.count("}")
            while depth > 0 and i + 1 < len(lines):
                i += 1
                title_buf += "\n" + lines[i]
                depth = title_buf.count("{") - title_buf.count("}")
            # 找 } 的位置
            arg = _balanced_arg("{" + title_buf, 0)
            if arg:
                flush_para()
                blocks.append(("title", arg[0]))
            i += 1
            continue
        if stripped in (r"\maketitle", r"\date{}"):
            i += 1
            continue
        if stripped.startswith(r"\linenumbers") or stripped.startswith(
            r"\renewcommand"
        ):
            i += 1
            continue

        # 标题：用平衡花括号提取标题文本（标题中可能含 $...$）
        heading_match = None
        for kind, cmd in (("h1", r"\section"), ("h2", r"\subsection"),
                          ("h3", r"\subsubsection")):
            # 允许 \section* 与 \section
            pattern = re.escape(cmd) + r"\*?\s*\{"
            m = re.match(pattern, stripped)
            if m:
                brace_start = stripped.find("{", m.end() - 1)
                arg = _balanced_arg(stripped, brace_start)
                if arg:
                    heading_match = (kind, arg[0])
                    break
        if heading_match:
            flush_para()
            blocks.append(heading_match)
            i += 1
            continue

        # 环境块：把 \begin{env} ... \end{env} 整段抽出
        m_env = re.match(r"\\begin\{(\w+\*?)\}", stripped)
        if m_env:
            env = m_env.group(1)
            flush_para()
            buf = [line]
            depth = 1
            i += 1
            while i < len(lines):
                buf.append(lines[i])
                if re.search(r"\\begin\{" + re.escape(env) + r"\}", lines[i]):
                    depth += 1
                if re.search(r"\\end\{" + re.escape(env) + r"\}", lines[i]):
                    depth -= 1
                    if depth == 0:
                        i += 1
                        break
                i += 1
            blocks.append(("env:" + env, "\n".join(buf)))
            # 若是 display-math 环境（equation/align），且其后紧跟非空行的
            # 正文（没有 LaTeX 段落分隔），则该正文应被视为公式所在段的
            # 延续，不再首行缩进——这是学术排版的通行约定。
            if env in ("equation", "equation*", "align", "align*"):
                if i < len(lines) and lines[i].strip() != "":
                    cont_flag[0] = True
            continue

        # 默认：累计为段落
        cur_para.append(line)
        i += 1

    flush_para()
    return blocks


def _emit_paragraph(doc: Document, text: str, *, indent_cm: float = 0.0,
                    first_line: bool = True):
    p = doc.add_paragraph()
    _format_paragraph(p, indent_cm=indent_cm, space_after_pt=SPACE_AFTER_PT,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      first_line_indent_cm=(FIRST_LINE_INDENT_CM
                                            if first_line else 0.0))
    # 段落里若有 \emph{... \eqref{X} ...} 嵌套，先把外层剥掉
    text = _flatten_nested_text_cmds(text)
    # 段内物理换行规整成空格；LaTeX 段落语义是空白换行才分段
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r" {2,}", " ", text).strip()
    _render_inline_segment(p, text)


def _emit_heading(doc: Document, level: int, text: str):
    # 用自定义段落而非 add_heading，保证 12pt 与字体一致
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT,
                      keep_with_next=True)
    # 编号前缀
    s = "S" if _S_PREFIX else ""
    if level == 1:
        _SECTION_COUNTERS["section"] += 1
        _SECTION_COUNTERS["subsection"] = 0
        _SECTION_COUNTERS["subsubsection"] = 0
        prefix = f"{s}{_SECTION_COUNTERS['section']}. "
        size = HEADING_PT + 2
    elif level == 2:
        _SECTION_COUNTERS["subsection"] += 1
        _SECTION_COUNTERS["subsubsection"] = 0
        prefix = (
            f"{s}{_SECTION_COUNTERS['section']}.{_SECTION_COUNTERS['subsection']} "
        )
        size = HEADING_PT + 1
    else:
        _SECTION_COUNTERS["subsubsection"] += 1
        prefix = (
            f"{s}{_SECTION_COUNTERS['section']}.{_SECTION_COUNTERS['subsection']}"
            f".{_SECTION_COUNTERS['subsubsection']} "
        )
        size = HEADING_PT
    run = p.add_run(prefix)
    _set_run_font(run, size_pt=size, bold=True)
    # 清理标题中的 \label{}
    clean = re.sub(r"\\label\{[^}]+\}", "", text)
    _render_inline_segment(p, clean)
    for r in p.runs:
        if r.font.bold is None:
            _set_run_font(r, size_pt=size, bold=True)
        else:
            r.font.size = Pt(size)
            r.font.bold = True


# ─── 环境处理器 ───────────────────────────────────────────────────


def _handle_abstract(doc: Document, body: str):
    inner = re.search(
        r"\\begin\{abstract\}(.*?)\\end\{abstract\}", body, re.DOTALL
    )
    if not inner:
        return
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT,
                      keep_with_next=True)
    run = p.add_run("Abstract")
    _set_run_font(run, size_pt=BODY_PT, bold=True)
    txt = inner.group(1)
    # 删除 \noindent
    txt = re.sub(r"\\noindent\b", "", txt)
    paragraphs = re.split(r"\n\s*\n", txt.strip())
    for para in paragraphs:
        if para.strip():
            _emit_paragraph(doc, para.strip())


def _handle_equation(doc: Document, body: str, *, env: str = "equation"):
    global _equation_counter
    # 提取内容（去掉 \begin{equation} ... \end{equation}）
    inner = re.search(
        r"\\begin\{" + env + r"\*?\}(.*?)\\end\{" + env + r"\*?\}",
        body,
        re.DOTALL,
    )
    if not inner:
        return
    eq_body = inner.group(1).strip()
    # 移除 \label{...} 并记录编号
    label_match = re.search(r"\\label\{([^}]+)\}", eq_body)
    eq_body = re.sub(r"\\label\{[^}]+\}", "", eq_body).strip()
    # 行末 LaTeX 标点（,/;/.）在 Word 公式里属于冗余，剥掉。
    eq_body = re.sub(r"\s*[,;.]\s*$", "", eq_body).strip()
    _equation_counter += 1
    eq_num = str(_equation_counter)
    expanded = _expand_macros(eq_body)
    for sp_cmd in [r'\,', r'\;', r'\:', r'\!']:
        expanded = expanded.replace(sp_cmd, ' ')
    p = doc.add_paragraph()
    # 左对齐 + 居中制表位 + 右对齐制表位：实现公式居中、编号右对齐
    _format_paragraph(p, space_after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(p._element, qn("w:pPr"))
    tabs = etree.SubElement(pPr, qn("w:tabs"))
    center_tab = etree.SubElement(tabs, qn("w:tab"))
    center_tab.set(qn("w:val"), "center")
    center_tab.set(qn("w:pos"), "4536")  # 页面中心 (9.07 cm / 2 ≈ 4536 twips)
    right_tab = etree.SubElement(tabs, qn("w:tab"))
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), "9072")  # 右边距
    # 制表符 → OMML 公式 → 制表符 → 编号
    pre_run = p.add_run("\t")
    _set_run_font(pre_run, size_pt=BODY_PT)
    omath = _latex_to_omath(expanded)
    if omath is not None:
        para_xml = etree.SubElement(p._element, qn("m:oMathPara"))
        para_xml.append(omath)
    else:
        run = p.add_run(eq_body)
        _set_run_font(run, size_pt=BODY_PT)
    tab_run = p.add_run("\t(" + eq_num + ")")
    _set_run_font(tab_run, size_pt=BODY_PT)


def _handle_align(doc: Document, body: str):
    """align/align* 当作多行 equation 处理，每个 \\\\ 一行。"""
    inner = re.search(
        r"\\begin\{align\*?\}(.*?)\\end\{align\*?\}", body, re.DOTALL
    )
    if not inner:
        return
    eq_body = inner.group(1)
    is_starred = "align*" in body
    # 按 \\ 分行
    rows = re.split(r"\\\\", eq_body)
    global _equation_counter
    for row in rows:
        row = row.strip()
        if not row:
            continue
        label_match = re.search(r"\\label\{([^}]+)\}", row)
        row = re.sub(r"\\label\{[^}]+\}", "", row).strip()
        row = row.replace("&", "")  # align 的对齐符号
        # 末尾标点同样剥掉
        row = re.sub(r"\s*[,;.]\s*$", "", row).strip()
        if not row:
            continue
        p = doc.add_paragraph()
        _format_paragraph(p, space_after_pt=4, align=WD_ALIGN_PARAGRAPH.CENTER)
        row_expanded = _expand_macros(row)
        for sp_cmd in [r'\,', r'\;', r'\:', r'\!']:
            row_expanded = row_expanded.replace(sp_cmd, ' ')
        omath = _latex_to_omath(row_expanded)
        if omath is not None:
            para_xml = etree.SubElement(p._element, qn("m:oMathPara"))
            para_xml.append(omath)
        else:
            run = p.add_run(row)
            _set_run_font(run)
        if not is_starred:
            _equation_counter += 1
            tab_run = p.add_run("    (" + str(_equation_counter) + ")")
            _set_run_font(tab_run)


def _handle_itemize(doc: Document, body: str, ordered: bool = False):
    """处理 itemize/enumerate。"""
    env = "enumerate" if ordered else "itemize"
    inner = re.search(
        r"\\begin\{" + env + r"\}(.*?)\\end\{" + env + r"\}", body, re.DOTALL
    )
    if not inner:
        return
    items = re.split(r"\\item\b", inner.group(1))
    counter = 0
    for item in items:
        item = item.strip()
        if not item:
            continue
        counter += 1
        marker = f"{counter}." if ordered else "•"
        p = doc.add_paragraph()
        _format_paragraph(p, indent_cm=0.8, space_after_pt=2,
                          align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run = p.add_run(marker + "  ")
        _set_run_font(run, bold=ordered)
        _render_inline_segment(p, item)


def _parse_table_rows(tabular_body: str) -> list[list[str]]:
    """解析 tabular 内部为二维列表。"""
    # 去掉 \toprule \midrule \bottomrule \hline
    body = tabular_body
    body = re.sub(r"\\toprule\b|\\midrule\b|\\bottomrule\b|\\hline\b", "", body)
    body = body.strip()
    rows = re.split(r"\\\\", body)
    parsed: list[list[str]] = []
    for row in rows:
        row = row.strip()
        if not row:
            continue
        # \multicolumn{N}{align}{content} → 把 content 放第一列，后面补 N-1 个空
        mcol_match = re.match(
            r"\\multicolumn\{(\d+)\}\{[^}]*\}\{(.*)\}\s*$", row, re.DOTALL
        )
        if mcol_match:
            n = int(mcol_match.group(1))
            content = mcol_match.group(2).strip()
            parsed.append([content] + [""] * (n - 1))
            continue
        # 按 & 分列（但不分裂 \& 转义）
        cells = re.split(r"(?<!\\)&", row)
        parsed.append([c.strip() for c in cells])
    return parsed


def _add_table_caption(doc: Document, label: str, caption: str):
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=4, align=WD_ALIGN_PARAGRAPH.LEFT,
                      keep_with_next=True)
    run = p.add_run(label + ". ")
    _set_run_font(run, bold=True)
    _render_inline_segment(p, caption)


def _extract_caption(body: str) -> str:
    """从 \\caption{...} 中用平衡花括号提取内容。"""
    m = re.search(r"\\caption\s*\{", body)
    if not m:
        return ""
    brace_start = body.find("{", m.end() - 1)
    arg = _balanced_arg(body, brace_start)
    return arg[0] if arg else ""


def _flatten_nested_text_cmds(s: str) -> str:
    """递归剥离 \\emph/\\textbf/\\textit/\\texttt 等嵌套包装。
    用于表格单元格等位置，让 inline 渲染器能处理只剩一层包装的内容。
    Math（$...$）内部不动。
    """
    changed = True
    iterations = 0
    while changed and iterations < 10:
        changed = False
        iterations += 1
        # 找最内层的 \cmd{plain content}（content 不含 { }）
        # 我们用平衡花括号定位，从左到右扫描
        out = []
        i = 0
        n = len(s)
        while i < n:
            m = re.match(r"\\(emph|textbf|textit|texttt|textsc|mathrm|mathbf)\s*\{",
                         s[i:])
            if m:
                cmd = m.group(1)
                brace_start = i + m.end() - 1
                arg = _balanced_arg(s, brace_start)
                if arg:
                    inner, end_pos = arg
                    # 若 inner 中还有 \emph/\textbf/\textit/\texttt 嵌套，
                    # 递归先剥外层（保留最外层标记一层）
                    if re.search(r"\\(emph|textbf|textit|texttt|textsc)\b", inner):
                        # 剥外层：直接放 inner（也许 inner 还需再剥）
                        out.append(inner)
                        changed = True
                        i = end_pos
                        continue
            out.append(s[i])
            i += 1
        s = "".join(out)
    return s


def _detect_midrule_rows(tabular_body: str) -> set[int]:
    """返回 \\midrule 后那一行的 0-based 索引集合（即"在这些行的上方画一条
    线"——三线表通常表头下需一条，分组之间也可能各一条）。
    """
    midrule_rows: set[int] = set()
    # 把 \\midrule 视作一个伪 token，扫描每个 \\ 分割段是否以 midrule 开头
    body = re.sub(r"\\toprule\b|\\bottomrule\b|\\hline\b", "", tabular_body).strip()
    pieces = re.split(r"\\\\", body)
    idx = 0
    for p in pieces:
        ps = p.strip()
        # 这一段在前一行的 \\ 之后；如果以 \midrule 开头，则它是 idx 行
        if not ps:
            continue
        if ps.startswith("\\midrule"):
            midrule_rows.add(idx)
            ps = re.sub(r"^\\midrule\s*", "", ps).strip()
            if not ps:
                continue
        idx += 1
    return midrule_rows


def _set_cell_border(cell, *, top=None, bottom=None,
                     left="nil", right="nil"):
    """单元格四边线控制。size 取 8 (=1pt) 模拟 booktabs 默认线宽。"""
    from docx.oxml.ns import qn as _qn
    tc = cell._tc
    tcPr = tc.find(_qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, _qn("w:tcPr"))
    # 移除既有 borders
    existing = tcPr.find(_qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = etree.SubElement(tcPr, _qn("w:tcBorders"))
    for side, spec in (("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)):
        el = etree.SubElement(borders, _qn("w:" + side))
        if spec == "nil" or spec is None:
            el.set(_qn("w:val"), "nil")
        else:
            # spec 可以是 ('single', size_in_eighth_pt) 或字符串 'single'
            if isinstance(spec, tuple):
                style, size = spec
            else:
                style, size = spec, 8
            el.set(_qn("w:val"), style)
            el.set(_qn("w:sz"), str(size))
            el.set(_qn("w:space"), "0")
            el.set(_qn("w:color"), "000000")


def _handle_table(doc: Document, body: str):
    """处理 table 环境（三线表 + 单元居中）。"""
    global _table_counter
    _table_counter += 1
    caption = _extract_caption(body)
    # 表格里的 caption 也要先剥嵌套
    caption = _flatten_nested_text_cmds(caption)
    # tabular
    tab_m = re.search(
        r"\\begin\{tabular\}\{[^}]*\}(.*?)\\end\{tabular\}", body, re.DOTALL
    )
    if not tab_m:
        return
    tabular_body = tab_m.group(1)
    rows = _parse_table_rows(tabular_body)
    if not rows:
        return
    midrule_rows = _detect_midrule_rows(tabular_body)
    s = "S" if _S_PREFIX else ""
    _add_table_caption(doc, f"Table {s}{_table_counter}", caption)
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    # 显式不要 Table Grid（避免任何竖线/水平线残留）
    table.style = "Normal Table"
    # 把整张表所有边线先关掉；后面只给指定行加 booktabs 横线
    from docx.oxml.ns import qn as _qn
    tblPr = table._tbl.find(_qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(table._tbl, _qn("w:tblPr"))
    old_bd = tblPr.find(_qn("w:tblBorders"))
    if old_bd is not None:
        tblPr.remove(old_bd)
    tbl_borders = etree.SubElement(tblPr, _qn("w:tblBorders"))
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(tbl_borders, _qn("w:" + s))
        el.set(_qn("w:val"), "nil")
    # 找哪些行是 multicolumn 全跨行
    raw_rows = re.split(r"\\\\",
        re.sub(r"\\(?:top|mid|bottom)rule\b|\\hline\b", "",
               tabular_body.strip()))
    full_span = []
    for raw in raw_rows:
        if not raw.strip():
            continue
        full_span.append(re.match(r"\s*\\multicolumn", raw.strip()) is not None)
    for i, row in enumerate(rows):
        spans = i < len(full_span) and full_span[i]
        # 先合并 multicolumn 行，再写内容
        if spans and ncols > 1:
            first_cell = table.rows[i].cells[0]
            for j in range(1, ncols):
                first_cell.merge(table.rows[i].cells[j])
        target_cols = 1 if spans else ncols
        for j in range(target_cols):
            cell = table.rows[i].cells[j]
            # 单元垂直居中
            tc = cell._tc
            tcPr = tc.find(_qn("w:tcPr"))
            if tcPr is None:
                tcPr = etree.SubElement(tc, _qn("w:tcPr"))
            old_vAlign = tcPr.find(_qn("w:vAlign"))
            if old_vAlign is not None:
                tcPr.remove(old_vAlign)
            vAlign = etree.SubElement(tcPr, _qn("w:vAlign"))
            vAlign.set(_qn("w:val"), "center")
            # 合并后 cell 内会有多段空 paragraph，全部删掉留 1 段
            paras = tc.findall(_qn("w:p"))
            for extra in paras[1:]:
                tc.remove(extra)
            p = cell.paragraphs[0]
            # 清空段内所有 run
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)
            # 单元水平居中——多列跨行（标题行）允许左对齐更自然，
            # 但常规数据行一律居中。
            if spans:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            content = row[j] if j < len(row) else ""
            content = _flatten_nested_text_cmds(content)
            _render_inline_segment(p, content)
            for run in p.runs:
                _set_run_font(run, size_pt=10.5, bold=(i == 0))
    # ── booktabs 三线 ──
    # 顶线：第 0 行所有 cell 顶边粗线 (size=12 ≈ 1.5pt)
    # 表头中线：i ∈ midrule_rows 的每一行顶边普通线 (size=6 ≈ 0.75pt)
    # 底线：最后一行所有 cell 底边粗线 (size=12)
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top=("single", 12))
    for mi in midrule_rows:
        if 0 <= mi < len(rows):
            for cell in table.rows[mi].cells:
                _set_cell_border(cell, top=("single", 6))
    for cell in table.rows[-1].cells:
        # 保留已有 top（首行同时是末行的退化情形也安全）
        tc = cell._tc
        tcPr = tc.find(_qn("w:tcPr"))
        existing = tcPr.find(_qn("w:tcBorders")) if tcPr is not None else None
        cur_top = None
        if existing is not None:
            top_el = existing.find(_qn("w:top"))
            if top_el is not None:
                val = top_el.get(_qn("w:val"))
                sz = top_el.get(_qn("w:sz"))
                if val and val != "nil":
                    cur_top = (val, int(sz) if sz else 8)
        _set_cell_border(cell, top=cur_top, bottom=("single", 12))


def _pdf_to_png(pdf_path: Path, dpi: int = 300) -> Optional[Path]:
    """用 pymupdf 把 PDF 第一页转 PNG，并写入 dpi metadata；返回缓存路径。"""
    try:
        import fitz  # type: ignore
    except ImportError:
        return None
    out = pdf_path.with_suffix(f".docx-{dpi}dpi.png")
    if out.exists() and out.stat().st_mtime >= pdf_path.stat().st_mtime:
        return out
    try:
        doc_pdf = fitz.open(str(pdf_path))
        page = doc_pdf.load_page(0)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        # 先存中间文件再用 PIL 重写以写入正确 DPI metadata
        tmp = out.with_suffix(".tmp.png")
        pix.save(str(tmp))
        doc_pdf.close()
        try:
            from PIL import Image  # type: ignore
            im = Image.open(tmp)
            im.save(str(out), dpi=(dpi, dpi))
            tmp.unlink(missing_ok=True)
        except Exception:
            tmp.rename(out)
        return out
    except Exception:
        return None


def _handle_figure(doc: Document, body: str):
    global _figure_counter
    _figure_counter += 1
    caption = _flatten_nested_text_cmds(_extract_caption(body))
    # 找图片路径
    img_m = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", body)
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=4, align=WD_ALIGN_PARAGRAPH.CENTER,
                      keep_with_next=True)
    if img_m:
        from docx.shared import Inches
        src = img_m.group(1).strip()
        # 试图找到图片文件
        src_path = Path(src)
        candidates = [
            src_path,
            src_path.with_suffix(".png"),
            src_path.with_suffix(".jpg"),
            src_path.with_suffix(".pdf"),
            Path("figures") / src,
            Path("figures") / (src_path.name + ".png"),
        ]
        embedded = False
        for cand in candidates:
            if not cand.is_file():
                continue
            suf = cand.suffix.lower()
            if suf in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
                try:
                    p.add_run().add_picture(str(cand), width=Inches(6.0))
                    embedded = True
                    break
                except Exception:
                    continue
            elif suf == ".pdf":
                # 现场转 PNG（300 dpi）
                png = _pdf_to_png(cand, dpi=300)
                if png and png.is_file():
                    try:
                        p.add_run().add_picture(str(png), width=Inches(6.0))
                        embedded = True
                        break
                    except Exception:
                        continue
        if not embedded:
            run = p.add_run(f"[Figure: {src}]")
            _set_run_font(run, italic=True)
    p_cap = doc.add_paragraph()
    _format_paragraph(p_cap, space_after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    s = "S" if _S_PREFIX else ""
    run = p_cap.add_run(f"Figure {s}{_figure_counter}. ")
    _set_run_font(run, bold=True)
    _render_inline_segment(p_cap, caption)


def _handle_algorithm(doc: Document, body: str):
    """algorithm + algorithmic：原样列出伪代码。"""
    caption = _flatten_nested_text_cmds(_extract_caption(body))
    lab_m = re.search(r"\\label\{([^}]+)\}", body)
    alg_id = _resolve_ref(lab_m.group(1)) if lab_m else "?"
    # 标题
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=4, align=WD_ALIGN_PARAGRAPH.LEFT,
                      keep_with_next=True)
    run = p.add_run(f"Algorithm {alg_id}. ")
    _set_run_font(run, bold=True)
    _render_inline_segment(p, caption)
    # algorithmic 块
    alg = re.search(
        r"\\begin\{algorithmic\}(?:\[\d+\])?(.*?)\\end\{algorithmic\}",
        body, re.DOTALL
    )
    if not alg:
        return
    raw_lines = [l for l in alg.group(1).split("\n") if l.strip()]
    line_no = 0
    # 合并续行：以 \State / \Require / \For / \If / \While 起的为新行；
    # 否则附加到上一行
    merged: list[str] = []
    for rl in raw_lines:
        s = rl.strip()
        if re.match(r"\\(State|Require|Ensure|For|EndFor|If|Else|EndIf|"
                    r"While|EndWhile|Return|Repeat|Until)\b", s):
            merged.append(s)
        else:
            if merged:
                merged[-1] += " " + s
            else:
                merged.append(s)
    indent = 0
    for s in merged:
        # 识别命令并提取
        comment_match = re.search(r"\\Comment\{([^}]*)\}", s)
        comment = comment_match.group(1) if comment_match else ""
        s_clean = re.sub(r"\\Comment\{[^}]*\}", "", s).strip()

        if s_clean.startswith(r"\Require"):
            text = "Require: " + s_clean[len(r"\Require"):].strip()
            cur_indent = 0
        elif s_clean.startswith(r"\Ensure"):
            text = "Ensure: " + s_clean[len(r"\Ensure"):].strip()
            cur_indent = 0
        elif s_clean.startswith(r"\For"):
            line_no += 1
            arg = re.match(r"\\For\s*\{(.+)\}", s_clean)
            inner = arg.group(1) if arg else s_clean[len(r"\For"):]
            text = f"{line_no}: for {inner} do"
            cur_indent = indent
            indent += 1
        elif s_clean.startswith(r"\EndFor"):
            indent = max(0, indent - 1)
            text = "end for"
            cur_indent = indent
        elif s_clean.startswith(r"\If"):
            line_no += 1
            arg = re.match(r"\\If\s*\{(.+)\}", s_clean)
            inner = arg.group(1) if arg else s_clean[len(r"\If"):]
            text = f"{line_no}: if {inner} then"
            cur_indent = indent
            indent += 1
        elif s_clean.startswith(r"\Else"):
            text = "else"
            cur_indent = max(0, indent - 1)
        elif s_clean.startswith(r"\EndIf"):
            indent = max(0, indent - 1)
            text = "end if"
            cur_indent = indent
        elif s_clean.startswith(r"\While"):
            line_no += 1
            arg = re.match(r"\\While\s*\{(.+)\}", s_clean)
            inner = arg.group(1) if arg else s_clean[len(r"\While"):]
            text = f"{line_no}: while {inner} do"
            cur_indent = indent
            indent += 1
        elif s_clean.startswith(r"\EndWhile"):
            indent = max(0, indent - 1)
            text = "end while"
            cur_indent = indent
        elif s_clean.startswith(r"\State"):
            line_no += 1
            inner = s_clean[len(r"\State"):].strip()
            # 进一步识别 \Return 嵌在 \State 中
            if inner.startswith(r"\Return"):
                rest = inner[len(r"\Return"):].strip()
                inner = "return" + (f" {rest}" if rest else "")
            text = f"{line_no}: " + inner
            cur_indent = indent
        elif s_clean.startswith(r"\Return"):
            line_no += 1
            rest = s_clean[len(r"\Return"):].strip()
            text = f"{line_no}: return" + (f" {rest}" if rest else "")
            cur_indent = indent
        else:
            text = s_clean
            cur_indent = indent
        # 写一行
        ap = doc.add_paragraph()
        _format_paragraph(ap, space_after_pt=0, align=WD_ALIGN_PARAGRAPH.LEFT)
        ap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        ap.paragraph_format.left_indent = Cm(1.0 + cur_indent * 0.6)
        _render_inline_segment(ap, text)
        if comment:
            run = ap.add_run("   ▷ ")
            _set_run_font(run, italic=True,
                          color=RGBColor(0x55, 0x55, 0x55))
            _render_inline_segment(ap, comment)
            # 让 comment 内容也斜体灰色
            for run in ap.runs[-2:]:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _handle_proposition(doc: Document, body: str):
    global _proposition_counter
    _proposition_counter += 1
    inner = re.search(
        r"\\begin\{proposition\}(?:\[([^\]]*)\])?(.*?)\\end\{proposition\}",
        body, re.DOTALL
    )
    if not inner:
        return
    name = inner.group(1)
    content = inner.group(2).strip()
    content = re.sub(r"\\label\{[^}]+\}", "", content).strip()
    content = _flatten_nested_text_cmds(content)
    p = doc.add_paragraph()
    _format_paragraph(p, space_after_pt=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(f"Proposition {_proposition_counter}")
    _set_run_font(run, bold=True)
    if name:
        run2 = p.add_run(f" ({name})")
        _set_run_font(run2, bold=True, italic=True)
    run3 = p.add_run(". ")
    _set_run_font(run3, bold=True)
    _render_inline_segment(p, content)


def _handle_thebibliography(doc: Document, body: str):
    """参考文献：按正文引用顺序输出 [1] [2] ... 编号。"""
    inner = re.search(
        r"\\begin\{thebibliography\}\{[^}]*\}(.*?)\\end\{thebibliography\}",
        body, re.DOTALL
    )
    if not inner:
        return
    _emit_heading(doc, 1, "References")
    # 收集所有 bibitem: key → text
    items: dict[str, str] = {}
    parts = re.split(r"\\bibitem\b", inner.group(1))
    for it in parts:
        it = it.strip()
        if not it:
            continue
        m = re.match(r"\{([^}]+)\}\s*(.*)", it, re.DOTALL)
        if m:
            items[m.group(1).strip()] = m.group(2).strip()
    # 按引用顺序输出
    for key in _cite_seq:
        text = items.get(key, "")
        if not text:
            # 未给出条目（理论上不该发生，因 _prescan_citations 已附加）
            text = f"[Reference key {key} not found in thebibliography.]"
        num = _cite_order[key]
        p = doc.add_paragraph()
        _format_paragraph(p, indent_cm=0.8, space_after_pt=4,
                          align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Cm(-0.8)
        run = p.add_run(f"[{num}] ")
        _set_run_font(run)
        text = re.sub(r"\s+", " ", text)
        _render_inline_segment(p, text)


# ─── 主转换函数 ─────────────────────────────────────────────────


ENV_DISPATCH = {
    "abstract": lambda d, b: _handle_abstract(d, b),
    "equation": lambda d, b: _handle_equation(d, b, env="equation"),
    "equation*": lambda d, b: _handle_equation(d, b, env="equation"),
    "align": lambda d, b: _handle_align(d, b),
    "align*": lambda d, b: _handle_align(d, b),
    "itemize": lambda d, b: _handle_itemize(d, b, ordered=False),
    "enumerate": lambda d, b: _handle_itemize(d, b, ordered=True),
    "table": lambda d, b: _handle_table(d, b),
    "figure": lambda d, b: _handle_figure(d, b),
    "algorithm": lambda d, b: _handle_algorithm(d, b),
    "proposition": lambda d, b: _handle_proposition(d, b),
    "thebibliography": lambda d, b: _handle_thebibliography(d, b),
}


def convert_tex_to_docx(tex_path: Path, out_path: Path):
    global _S_PREFIX
    _S_PREFIX = "supplementary" in str(tex_path).lower()
    source = tex_path.read_text(encoding="utf-8")
    # 去除 BOM 和 CRLF
    source = source.replace("\r\n", "\n").replace("\r", "\n")
    source = _normalize_source(source)
    _prescan_labels(source)
    _prescan_citations(source)

    doc = Document()
    # 全局样式：12pt Times New Roman, double-spaced
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(BODY_PT)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), DEFAULT_FONT)
    rfonts.set(qn("w:hAnsi"), DEFAULT_FONT)
    rfonts.set(qn("w:cs"), DEFAULT_FONT)

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    blocks = _split_logical_blocks(source)

    for kind, payload in blocks:
        if kind == "title":
            title = re.sub(r"\\textbf\{(.+?)\}", r"\1", payload, flags=re.DOTALL)
            title = re.sub(r"\s+", " ", title).strip()
            p = doc.add_paragraph()
            _format_paragraph(p, space_after_pt=12,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(title)
            _set_run_font(run, size_pt=14, bold=True)
        elif kind == "h1":
            _emit_heading(doc, 1, payload)
        elif kind == "h2":
            _emit_heading(doc, 2, payload)
        elif kind == "h3":
            _emit_heading(doc, 3, payload)
        elif kind == "para":
            # \noindent\textbf{Keywords:} 这类段落
            _emit_paragraph(doc, payload)
        elif kind == "para_cont":
            # 紧跟 display-math 的延续段落：不加首行缩进。
            _emit_paragraph(doc, payload, first_line=False)
        elif kind.startswith("env:"):
            env = kind[4:]
            handler = ENV_DISPATCH.get(env)
            if handler:
                handler(doc, payload)
            else:
                # 未知环境：尝试剥外壳后当段落处理
                inner = re.sub(
                    r"\\begin\{" + re.escape(env) + r"\}|"
                    r"\\end\{" + re.escape(env) + r"\}", "", payload
                )
                if inner.strip():
                    _emit_paragraph(doc, inner.strip())

    doc.save(str(out_path))
    print(f"Written: {out_path}")
    print(f"  Sections: {_SECTION_COUNTERS['section']}")
    print(f"  Equations: {_equation_counter}")
    print(f"  Tables: {_table_counter}")
    print(f"  Figures: {_figure_counter}")
    print(f"  Propositions: {_proposition_counter}")
    print(f"  Labels mapped: {len(_label_map)}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-i", "--input", required=True)
    ap.add_argument("-o", "--output", required=True)
    args = ap.parse_args()
    convert_tex_to_docx(Path(args.input).resolve(),
                        Path(args.output).resolve())


if __name__ == "__main__":
    main()
